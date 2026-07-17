"""将 Word 练习题文档切分为可进入题库流程的原始 JSON 习题。"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any


QUESTION_START = re.compile(
    r"^(?:\d+[.、．]|Given\b|Define(?!\s+(?:a\s+)?(?:derived|another)\b)\b|Please\b|Write\b|Implement\b|In the following\b|According to\b|A stack\b|Important aspects\b|选择|打开文件|假设|对下面|阅读下面|用C\+\+|已知类|下表中)",
    flags=re.IGNORECASE,
)
SECTION_START = re.compile(r"^第[一二三四五六七八九十\d]+篇")
ANSWER_SECTION = re.compile(r"^第3篇|常见错误", flags=re.IGNORECASE)
CODE_HINT = re.compile(r"(^#include\b|[{};]|\b(class|struct|template|namespace|void|int|double|char|bool|public|private|protected|virtual|return|cout|cin|main)\b|//)")


def read_docx(path: Path) -> tuple[list[str], list[list[str]]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法读取 DOCX。") from exc
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = [[cell.text.strip() for row in table.rows for cell in row.cells if cell.text.strip()] for table in document.tables]
    return paragraphs, tables


def split_question_blocks(paragraphs: list[str]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    current: list[str] = []
    section = ""
    collecting = False
    start_index = 1
    for index, paragraph in enumerate(paragraphs, start=1):
        if SECTION_START.match(paragraph):
            if current:
                blocks.append({"section": section, "start": start_index, "end": index - 1, "paragraphs": current})
                current = []
            section = paragraph
            collecting = not bool(ANSWER_SECTION.search(section))
            start_index = index
            continue
        if not collecting:
            continue
        if QUESTION_START.match(paragraph) and current:
            blocks.append({"section": section, "start": start_index, "end": index - 1, "paragraphs": current})
            current = []
            start_index = index
        if not current:
            start_index = index
        current.append(paragraph)
    if current and collecting:
        blocks.append({"section": section, "start": start_index, "end": len(paragraphs), "paragraphs": current})
    return [block for block in blocks if is_question_like(block["paragraphs"])]


def is_question_like(paragraphs: list[str]) -> bool:
    text = "\n".join(paragraphs).lower()
    if text.startswith("打开文件"):
        # 文档只引用了另一份 PDF，当前没有题干，不能把它当成可自动作答的习题。
        return False
    markers = ["class", "struct", "program", "代码", "函数", "输出", "定义", "实现", "错误", "队列", "栈", "继承", "template"]
    first_line = paragraphs[0] if paragraphs else ""
    has_question_signal = bool(QUESTION_START.match(first_line)) or any("?" in line or "？" in line for line in paragraphs)
    return len(text) >= 30 and has_question_signal and any(marker in text for marker in markers)


def classify_question(paragraphs: list[str]) -> str:
    text = "\n".join(paragraphs).lower()
    if any(marker in text for marker in ["错误所在", "改正", "改错", "error"]):
        return "code_fixing"
    if any(marker in text for marker in ["输出", "write the output", "which type-conversions"]):
        return "code_reading"
    if any(marker in text for marker in ["define", "implement", "实现", "定义", "design"]):
        return "short_programming"
    return "code_reading" if any(CODE_HINT.search(line) for line in paragraphs) else "short_programming"


def split_stem_and_code(paragraphs: list[str]) -> tuple[str, str]:
    stem_lines: list[str] = []
    code_lines: list[str] = []
    code_mode = False
    for line in paragraphs:
        # 题干中也常出现 class、函数、代码等单词；只有明显像源代码的整行才开启代码区。
        is_code = bool(
            re.match(
                r"^(?:#include\b|(?:class|struct)\s+[A-Za-z_]\w*.*[{:]|(?:public|private|protected)\s*:|(?:void|int|double|float|char|bool|auto|[A-Za-z_]\w*\s*[*&]?)\s+[A-Za-z_]\w*\s*\([^)]*\)\s*\{?|[{}]|//)",
                line.strip(),
            )
        )
        if is_code:
            code_mode = True
        if code_mode and (is_code or line.startswith((" ", "\t")) or line in {"}", "{"}):
            code_lines.append(line)
        else:
            stem_lines.append(line)
    stem = "\n".join(stem_lines).strip()
    code = "\n".join(code_lines).strip()
    if not stem:
        stem = paragraphs[0] if paragraphs else ""
    return stem, code


def infer_difficulty(question_type: str, paragraphs: list[str]) -> int:
    text = "\n".join(paragraphs).lower()
    if any(marker in text for marker in ["template", "虚函数", "多重", "类型转换", "class hierarchy"]):
        return 4
    if question_type in {"code_fixing", "code_reading"}:
        return 3
    return 3


def make_question(question_id: str, block: dict[str, Any], source_file: str) -> dict[str, Any]:
    paragraphs = block["paragraphs"]
    stem, code = split_stem_and_code(paragraphs)
    question_type = classify_question(paragraphs)
    return {
        "question_id": question_id,
        "type": question_type,
        "language": "C++",
        "stem": stem,
        "code": code,
        "options": [],
        "answer": "",
        "analysis": "",
        "answer_source": "missing",
        "answer_kind": "unknown",
        "answer_status": "needs_generation",
        "answer_confidence": 0.0,
        "difficulty": infer_difficulty(question_type, paragraphs),
        "abilities": infer_abilities(question_type),
        "gold_knowledge_points": [],
        "source": {
            "kind": "docx_exercise",
            "file": source_file,
            "section": block["section"],
            "paragraph_start": block["start"],
            "paragraph_end": block["end"],
            "parser": "docx_question_parser_v1",
        },
    }


def infer_abilities(question_type: str) -> list[str]:
    if question_type == "code_fixing":
        return ["错误定位", "代码修正"]
    if question_type == "code_reading":
        return ["代码阅读", "运行结果分析"]
    return ["程序设计", "代码实现"]


def table_questions(tables: list[list[str]], source_file: str, start_index: int) -> list[dict[str, Any]]:
    questions: list[dict[str, Any]] = []
    for table_index, cells in enumerate(tables, start=1):
        for cell_index, code in enumerate(cells, start=1):
            if len(code) < 20:
                continue
            question_id = f"CPP_DOCX_{start_index + len(questions):03d}"
            questions.append(
                {
                    "question_id": question_id,
                    "type": "code_fixing",
                    "language": "C++",
                    "stem": "指出下面代码片段中的错误，并说明原因。",
                    "code": code,
                    "options": [],
                    "answer": "",
                    "analysis": "",
                    "answer_source": "missing",
                    "answer_kind": "unknown",
                    "answer_status": "needs_generation",
                    "answer_confidence": 0.0,
                    "difficulty": 3,
                    "abilities": ["错误定位", "代码修正"],
                    "gold_knowledge_points": [],
                    "source": {"kind": "docx_exercise_table", "file": source_file, "table": table_index, "cell": cell_index, "parser": "docx_question_parser_v1"},
                }
            )
    return questions


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="将 Word 练习题解析为结构化 JSON。")
    parser.add_argument("--input", required=True, help="DOCX 习题文档路径")
    parser.add_argument("--output", default="work/oop_kg_demo/output/programming_kg/questions/docx_questions_raw.json", help="原始习题 JSON 输出路径")
    parser.add_argument("--report", default="work/oop_kg_demo/output/programming_kg/questions/docx_question_parse_report.json", help="解析报告输出路径")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    input_path = Path(args.input)
    paragraphs, tables = read_docx(input_path)
    blocks = split_question_blocks(paragraphs)
    questions = [make_question(f"CPP_DOCX_{index:03d}", block, input_path.name) for index, block in enumerate(blocks, start=1)]
    questions.extend(table_questions(tables, input_path.name, len(questions) + 1))
    output_path = Path(args.output)
    report_path = Path(args.report)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(questions, ensure_ascii=False, indent=2), encoding="utf-8")
    report = {
        "schema_version": "programming_kg_docx_question_parse_v1",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "input": str(input_path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "question_count": len(questions),
        "answer_provided_count": 0,
        "answer_generation_required_count": len(questions),
        "output": str(output_path),
    }
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"解析习题数量：{len(questions)}")
    print(f"原始习题：{output_path}")
    print(f"解析报告：{report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
